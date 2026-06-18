import io, math
from datetime import datetime
import pandas as pd
import numpy as np
import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.tri as mtri
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.ticker import FuncFormatter, MultipleLocator

# === PDF profesional con ReportLab ===
import os
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# === Informe Word editable ===
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


st.set_page_config(page_title="Planilla de inspección de represas", layout="wide")
st.title("Planilla de inspección de represas")
st.write("Carga puntos topográficos, calcula parámetros geométricos y genera un informe PDF estilo planilla de inspección.")

uploaded = st.file_uploader("Subí tu archivo Excel/CSV/TXT", type=["csv", "txt", "xlsx"])

def read_file(file):
    name = file.name.lower()
    if name.endswith(".xlsx"):
        return pd.read_excel(file)
    if name.endswith(".txt"):
        content = file.getvalue().decode("utf-8", errors="ignore")
        try:
            return pd.read_csv(io.StringIO(content), sep=None, engine="python")
        except Exception:
            return pd.read_csv(io.StringIO(content), sep=r"\s+", header=None)
    return pd.read_csv(file, sep=None, engine="python")

def norm_text(x):
    return str(x).strip().lower()


def parse_float_manual(valor):
    """Convierte texto manual a float. Acepta coma o punto decimal."""
    if valor is None:
        return None
    txt = str(valor).strip().replace(",", ".")
    if txt == "":
        return None
    try:
        return float(txt)
    except Exception:
        return None


def format_pk(m):
    m = float(m)
    km = int(m // 1000)
    rest = m - km * 1000
    return f"{km}+{rest:06.2f}"

def pk_formatter(x, pos):
    return format_pk(x)

def calcular_distancia_acumulada(df):
    out = df.copy().reset_index(drop=True)
    out["Distancia tramo"] = np.sqrt(out["X"].diff()**2 + out["Y"].diff()**2).fillna(0)
    out["Distancia acumulada"] = out["Distancia tramo"].cumsum()
    out["Progresiva"] = out["Distancia acumulada"].apply(format_pk)
    return out

def calcular_anchos_pares(base_df):
    base_df = base_df.copy()
    base_df["Numero_num"] = pd.to_numeric(base_df["Numero"], errors="coerce")
    base_df = base_df.sort_values("Numero_num").reset_index(drop=True)
    registros = []
    for i in range(0, len(base_df) - 1, 2):
        p1 = base_df.iloc[i]
        p2 = base_df.iloc[i + 1]
        H = math.sqrt((p2["X"] - p1["X"])**2 + (p2["Y"] - p1["Y"])**2)
        registros.append({
            "Base 1": p1["Numero"],
            "Base 2": p2["Numero"],
            "Cota Base 1": p1["Z"],
            "Cota Base 2": p2["Z"],
            "Ancho corona (m)": H,
            "Diferencia cota (m)": p2["Z"] - p1["Z"]
        })
    return pd.DataFrame(registros)

def create_mdt_figure(data, cmap="terrain", mostrar_curvas=True, intervalo=0.2, curvas_etiquetas=True, for_pdf=False):
    x, y, z = data["X"].to_numpy(), data["Y"].to_numpy(), data["Z"].to_numpy()
    triang = mtri.Triangulation(x, y)
    fig, ax = plt.subplots(figsize=(8.3, 7.3) if for_pdf else (10, 10), dpi=140)
    mdt = ax.tripcolor(triang, z, shading="gouraud", cmap=cmap)

    if mostrar_curvas:
        start = math.floor(float(np.nanmin(z)) / intervalo) * intervalo
        end = math.ceil(float(np.nanmax(z)) / intervalo) * intervalo
        levels = np.arange(start, end + intervalo, intervalo)
        if len(levels) > 1:
            cs = ax.tricontour(triang, z, levels=levels, linewidths=0.8)
            if curvas_etiquetas:
                ax.clabel(cs, inline=True, fontsize=7, fmt="%.2f")

    ax.scatter(x, y, s=12)
    ax.set_title("Modelo Digital del Terreno", fontsize=13, fontweight="bold")
    ax.set_xlabel("X / Este")
    ax.set_ylabel("Y / Norte")
    ax.set_aspect("equal", adjustable="box")
    ax.grid(alpha=0.25)
    fig.colorbar(mdt, ax=ax, label="Cota (m)")
    fig.tight_layout()
    return fig

def create_profile_figure(
    perfil,
    title_extra="",
    for_pdf=False,
    y_min_manual=None,
    y_max_manual=None,
    mostrar_marcadores=True,
    sombreado=True,
    mostrar_min=True,
    mostrar_max=True,
    min_dx=15,
    min_dy=-35,
    max_dx=-80,
    max_dy=25
):
    fig, ax = plt.subplots(figsize=(11.69, 5.2) if for_pdf else (15, 6), dpi=140)
    perfil = calcular_distancia_acumulada(perfil)
    x = perfil["Distancia acumulada"].to_numpy()
    y = perfil["Z"].to_numpy()

    ymin_default = math.floor(float(np.nanmin(y))*2)/2
    ymax_default = math.ceil(float(np.nanmax(y))*2)/2
    ymin = y_min_manual if y_min_manual is not None else ymin_default
    ymax = y_max_manual if y_max_manual is not None else ymax_default

    ax.set_facecolor("#FAFAFA")

    if sombreado:
        ax.fill_between(x, y, ymin, alpha=0.14)

    marker = "o" if mostrar_marcadores else None
    ax.plot(x, y, linewidth=2.4, marker=marker, markersize=4.5, label="Perfil relevado")
    ax.set_ylim(ymin, ymax)

    ax.xaxis.set_major_formatter(FuncFormatter(pk_formatter))
    ax.xaxis.set_major_locator(MultipleLocator(50))
    ax.xaxis.set_minor_locator(MultipleLocator(10))
    ax.yaxis.set_major_locator(MultipleLocator(0.5))
    ax.yaxis.set_minor_locator(MultipleLocator(0.1))
    ax.grid(which="major", linewidth=0.8, alpha=0.55)
    ax.grid(which="minor", linewidth=0.35, alpha=0.25)
    ax.set_title(f"Perfil longitudinal {title_extra}\nLongitud: {x[-1]:.2f} m", fontsize=13, fontweight="bold")
    ax.set_xlabel("Progresiva / Distancia acumulada (m)")
    ax.set_ylabel("Cota (m)")

    idx_min = int(np.nanargmin(y))
    idx_max = int(np.nanargmax(y))

    if mostrar_min:
        ax.scatter([x[idx_min]], [y[idx_min]], s=60, zorder=5)
        ax.annotate(
            f"Mín: {y[idx_min]:.3f} m\nPK {format_pk(x[idx_min])}",
            xy=(x[idx_min], y[idx_min]),
            xytext=(min_dx, min_dy),
            textcoords="offset points",
            fontsize=8,
            bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="0.65"),
            arrowprops=dict(arrowstyle="->", lw=0.8)
        )

    if mostrar_max:
        ax.scatter([x[idx_max]], [y[idx_max]], s=60, zorder=5)
        ax.annotate(
            f"Máx: {y[idx_max]:.3f} m\nPK {format_pk(x[idx_max])}",
            xy=(x[idx_max], y[idx_max]),
            xytext=(max_dx, max_dy),
            textcoords="offset points",
            fontsize=8,
            bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="0.65"),
            arrowprops=dict(arrowstyle="->", lw=0.8)
        )

    ax.legend()
    fig.tight_layout()
    return fig, perfil
def plot_talud(H, V, punto_sup, punto_inf, relacion, for_pdf=False):
    fig, ax = plt.subplots(figsize=(7.8, 4.2) if for_pdf else (7.2, 4.3), dpi=140)
    ax.set_facecolor("#FAFAFA")
    ax.plot([0, H], [V, 0], linewidth=2.5, marker="o")
    ax.plot([0, 0], [0, V], linestyle="--", linewidth=1.2)
    ax.plot([0, H], [0, 0], linestyle="--", linewidth=1.2)
    ax.text(0, V, f"  Superior\n  Pto {punto_sup}", va="bottom", fontsize=9)
    ax.text(H, 0, f"  Inferior\n  Pto {punto_inf}", va="top", fontsize=9)
    ax.text(H/2, -0.08*max(V,1), f"H = {H:.3f} m", ha="center", va="top")
    ax.text(-0.05*max(H,1), V/2, f"V = {V:.3f} m", ha="right", va="center", rotation=90)
    ax.set_title(f"Talud aguas abajo = 1V : {relacion:.2f}H", fontsize=13, fontweight="bold")
    ax.set_xlabel("Horizontal (m)")
    ax.set_ylabel("Vertical (m)")
    ax.set_aspect("equal", adjustable="box")
    ax.grid(alpha=0.25)
    ax.set_xlim(-0.15*max(H,1), H*1.15)
    ax.set_ylim(-0.20*max(V,1), V*1.25)
    fig.tight_layout()
    return fig

def add_header(ax, title):
    ax.axis("off")
    ax.text(0.03, 0.96, title, fontsize=18, fontweight="bold", transform=ax.transAxes, va="top")
    ax.plot([0.03, 0.97], [0.91, 0.91], transform=ax.transAxes, linewidth=1.2)

def text_lines_page(title, lines):
    fig, ax = plt.subplots(figsize=(8.27, 11.69), dpi=140)
    add_header(ax, title)
    y = 0.86
    for line in lines:
        if y < 0.08:
            break
        ax.text(0.06, y, line, fontsize=10.5, transform=ax.transAxes, va="top")
        y -= 0.04 if line else 0.025
    fig.tight_layout()
    return fig

def table_page(title, df_table, max_rows=32):
    fig, ax = plt.subplots(figsize=(11.69, 8.27), dpi=140)
    ax.axis("off")
    ax.set_title(title, fontsize=16, fontweight="bold", loc="left", pad=18)
    show = df_table.head(max_rows).copy()
    table = ax.table(cellText=show.values, colLabels=show.columns, loc="center", cellLoc="center")
    table.auto_set_font_size(False)
    table.set_fontsize(7.5)
    table.scale(1, 1.25)
    fig.tight_layout()
    return fig

def make_planilla_lines(form, auto):
    return [
        "PLANILLA INSPECCIÓN",
        "",
        "1. DATOS GENERALES",
        f"1.1 N° de represa: {form['n_represa']}",
        f"1.2 Fecha inspección: {form['fecha']}",
        f"1.3 Curso embalsado: {form['curso']}",
        f"1.4 Nombre del propietario: {form['propietario']}",
        f"1.5 Padrones represa: {form['padrones']}",
        f"    Sección Judicial: {form['seccion']}",
        f"    Departamento: {form['departamento']}",
        "",
        "2. DATOS REPRESA",
        f"2.1 Ancho mínimo de coronamiento: {auto.get('ancho_min', 'S/D')}",
        f"2.2 Sección de máxima altura:",
        f"    Altura del terraplén: {auto.get('altura_max', 'S/D')}",
        f"    Ancho de coronamiento: {auto.get('ancho_prom', 'S/D')}",
        "",
        "2.3 TALUD AGUAS ARRIBA",
        f"    {form['talud_arriba']} Pendiente calculada: {auto.get('talud_arriba_calc', 'S/D')}",
        "",
        "2.4 TALUD AGUAS ABAJO",
        f"    {form['talud_abajo']} Pendiente calculada: {auto.get('talud', 'S/D')}",
        "",
        f"2.5 Cota del Pelo de agua: {form['pelo_agua']}",
        f"2.6 Cota del Coronamiento: {auto.get('cota_corona', 'S/D')}",
        f"2.7 Escala: {form['escala']}",
        f"    Tipo y ubicación aproximada: {form['escala_ubicacion']}",
        f"2.8 Mojón de referencia: {form['mojon']}",
        f"    Ubicación aproximada: {form['mojon_ubicacion']}",
        f"    Cota: {form['mojon_cota']}",
        f"2.9 Origen utilizado en la nivelación: {form['origen']}",
    ]

def make_planilla_lines_2(form, auto):
    return [
        "2.10 OBRAS DE TOMA",
        f"Ubicación: {form['toma_ubicacion']}",
        f"Tipo y diámetro: {form['toma_tipo']}",
        f"Cota de zampeado aguas abajo: {form['toma_zampeado']}",
        "",
        "2.11 SEGURIDAD DE LA REPRESA",
        f"Situación observada hace temer por la seguridad de la obra: {form['seguridad']}",
        f"Comentarios: {form['seguridad_comentarios']}",
        "",
        "2.12 VERTEDERO",
        f"Franquía: {auto.get('franquia', 'S/D')}",
        f"Cota mínima corona utilizada: {auto.get('cota_min_corona_franquia', 'S/D')}",
        f"Cota mínima vertedero utilizada: {auto.get('cota_min_vertedero_franquia', 'S/D')}",
        f"Ubicación: {form['vertedero_ubicacion']}",
        f"Ancho del vertedero: {auto.get('ancho_vertedero', 'S/D')}",
        f"Cotas del vertedero: {auto.get('cotas_vertedero', 'S/D')}",
        f"Comentarios: {form['vertedero_comentarios']}",
        "",
        "2.13 OBSERVACIONES",
        form["observaciones"],
        "",
        "2.14 Forma de acceso a la represa",
        form["acceso"],
        "",
        f"2.15 Inspeccionada por: {form['inspectores']}",
    ]

def build_pdf(form, auto, data, fig_mdt, fig_profile, anchos_df, talud_info, fig_talud, perfil_table, incluir_talud_abajo=True, talud_arriba_info=None, fig_talud_arriba=None, incluir_talud_arriba=False):
    buffer = io.BytesIO()
    with PdfPages(buffer) as pdf:
        pdf.savefig(text_lines_page("Planilla de inspección", make_planilla_lines(form, auto)), bbox_inches="tight")
        plt.close()
        pdf.savefig(text_lines_page("Planilla de inspección", make_planilla_lines_2(form, auto)), bbox_inches="tight")
        plt.close()

        pdf.savefig(fig_mdt, bbox_inches="tight")
        plt.close(fig_mdt)
        pdf.savefig(fig_profile, bbox_inches="tight")
        plt.close(fig_profile)

        if anchos_df is not None and len(anchos_df):
            resumen_anchos = pd.DataFrame({
                "Indicador": ["Anchos calculados", "Ancho mínimo", "Ancho máximo", "Ancho promedio"],
                "Valor": [
                    len(anchos_df),
                    f"{anchos_df['Ancho corona (m)'].min():.3f} m",
                    f"{anchos_df['Ancho corona (m)'].max():.3f} m",
                    f"{anchos_df['Ancho corona (m)'].mean():.3f} m",
                ]
            })
            pdf.savefig(table_page("Resumen de anchos de coronamiento", resumen_anchos), bbox_inches="tight")
            plt.close()
            pdf.savefig(table_page("Tabla de anchos de coronamiento", anchos_df.round(3)), bbox_inches="tight")
            plt.close()

        if talud_info:
            talud_df = pd.DataFrame({
                "Parámetro": ["Punto superior", "Punto inferior", "Distancia horizontal H", "Diferencia vertical V", "Relación"],
                "Valor": [talud_info["p_sup"], talud_info["p_inf"], f"{talud_info['H']:.3f} m", f"{talud_info['V']:.3f} m", f"1V : {talud_info['relacion']:.2f}H"]
            })
            pdf.savefig(table_page("Talud aguas abajo", talud_df), bbox_inches="tight")
            plt.close()
            pdf.savefig(fig_talud, bbox_inches="tight")
            plt.close(fig_talud)

        if perfil_table is not None and len(perfil_table):
            pdf.savefig(table_page("Tabla del perfil longitudinal", perfil_table.round(3)), bbox_inches="tight")
            plt.close()

    buffer.seek(0)
    return buffer


# ================================
# PDF PROFESIONAL - VERSION MEJORADA
# Esta funcion reemplaza el PDF viejo sin tocar la parte visual de la app.
# ================================

def _registrar_verdana():
    posibles = [
        r"C:\Windows\Fonts\verdana.ttf",
        r"C:\Windows\Fonts\Verdana.ttf",
    ]
    for fuente in posibles:
        if os.path.exists(fuente):
            try:
                pdfmetrics.registerFont(TTFont("Verdana", fuente))
                return "Verdana"
            except Exception:
                pass
    return "Helvetica"

PDF_FONT = _registrar_verdana()

def _safe_text(x):
    if x is None:
        return ""
    txt = str(x)
    return txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")

def _guardar_figura_pdf(fig, nombre):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f"_{nombre}.png")
    fig.savefig(tmp.name, dpi=230, bbox_inches="tight")
    plt.close(fig)
    return tmp.name

def _styles_pdf():
    ss = getSampleStyleSheet()

    ss.add(ParagraphStyle(
        name="TituloPlanilla",
        fontName=PDF_FONT,
        fontSize=13,
        leading=16,
        alignment=1,
        spaceAfter=12,
    ))

    ss.add(ParagraphStyle(
        name="Seccion",
        fontName=PDF_FONT,
        fontSize=12,
        leading=15,
        spaceBefore=8,
        spaceAfter=6,
    ))

    ss.add(ParagraphStyle(
        name="Texto",
        fontName=PDF_FONT,
        fontSize=11,
        leading=14,
        spaceAfter=4,
    ))

    ss.add(ParagraphStyle(
        name="TextoChico",
        fontName=PDF_FONT,
        fontSize=9,
        leading=11,
        spaceAfter=3,
    ))

    ss.add(ParagraphStyle(
        name="Caption",
        fontName=PDF_FONT,
        fontSize=9,
        leading=11,
        alignment=1,
        textColor=colors.HexColor("#444444"),
        spaceBefore=3,
        spaceAfter=6,
    ))

    ss.add(ParagraphStyle(
        name="NotaTabla",
        fontName=PDF_FONT,
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#444444"),
        leftIndent=0.15 * cm,
        spaceBefore=4,
        spaceAfter=6,
    ))

    return ss

def _p(texto, estilo):
    return Paragraph(_safe_text(texto), estilo)

def _pb(texto, estilo):
    return Paragraph(f"<b>{_safe_text(texto)}</b>", estilo)

def _linea(label, valor, ss):
    return Paragraph(f"<b>{_safe_text(label)}:</b> {_safe_text(valor)}", ss["Texto"])

def _tabla_pdf(df, font_size=8, header_color="#D9D9D9"):
    data_tbl = [list(df.columns)] + df.astype(str).values.tolist()
    tabla = Table(data_tbl, repeatRows=1)
    tabla.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), PDF_FONT),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_color)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#8A8A8A")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tabla

def build_pdf(form, auto, data, fig_mdt, fig_profile, anchos_df, talud_info, fig_talud, perfil_table, incluir_talud_abajo=True, talud_arriba_info=None, fig_talud_arriba=None, incluir_talud_arriba=False):
    """
    Informe PDF profesional:
    - A4 vertical
    - margenes uniformes
    - fuente Verdana si esta disponible
    - secciones estilo planilla
    - graficos en paginas separadas
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.1 * cm,
        rightMargin=2.0 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title="Planilla de Inspeccion de Represa",
    )

    ss = _styles_pdf()
    story = []

    # Imagenes temporales
    mdt_path = _guardar_figura_pdf(fig_mdt, "mdt")
    profile_path = _guardar_figura_pdf(fig_profile, "perfil")
    talud_path = _guardar_figura_pdf(fig_talud, "talud") if fig_talud is not None else None
    talud_arriba_path = _guardar_figura_pdf(fig_talud_arriba, "talud_arriba") if fig_talud_arriba is not None else None

    # -------------------------
    # Pagina 1
    # -------------------------
    story.append(_pb("PLANILLA INSPECCION", ss["TituloPlanilla"]))

    story.append(_pb("1. DATOS GENERALES", ss["Seccion"]))
    story.append(_linea("1.1 N° de represa (segun listado Depto. Adm. Aguas)", form.get("n_represa", ""), ss))
    story.append(_linea("1.2 Fecha inspeccion", form.get("fecha", ""), ss))
    story.append(_linea("1.3 Curso embalsado", form.get("curso", ""), ss))
    story.append(_linea("1.4 Nombre del propietario", form.get("propietario", ""), ss))
    story.append(_linea("1.5 Padrones represa", form.get("padrones", ""), ss))
    story.append(_linea("Seccion Judicial", form.get("seccion", ""), ss))
    story.append(_linea("Departamento", form.get("departamento", ""), ss))

    story.append(Spacer(1, 0.15 * cm))
    story.append(_pb("2. DATOS REPRESA", ss["Seccion"]))
    story.append(_linea("2.1 Ancho minimo de coronamiento", auto.get("ancho_min", "S/D"), ss))

    story.append(_p("<b>2.2 Seccion de maxima altura</b>", ss["Texto"]))
    story.append(_p(f": Altura del terraplen: {auto.get('altura_max', 'S/D')}", ss["Texto"]))
    story.append(_p(f": Ancho de coronamiento: {auto.get('ancho_prom', 'S/D')}", ss["Texto"]))

    story.append(_pb("2.3 TALUD AGUAS ARRIBA", ss["Seccion"]))
    story.append(_p("(Describir estado actual; pendiente, enrocado, erosiones)", ss["Texto"]))
    story.append(_p(f": {form.get('talud_arriba', '')} Pendiente calculada = {auto.get('talud_arriba_calc', 'S/D')}", ss["Texto"]))

    story.append(_pb("2.4 TALUD AGUAS ABAJO", ss["Seccion"]))
    story.append(_p("(Describir estado actual; pendiente, filtraciones, empastado, erosiones)", ss["Texto"]))
    story.append(_p(f": {form.get('talud_abajo', '')} Pendiente calculada = {auto.get('talud', 'S/D')}", ss["Texto"]))

    story.append(_linea("2.5 Cota del Pelo de agua", form.get("pelo_agua", ""), ss))
    story.append(_linea("2.6 Cota del Coronamiento", auto.get("cota_corona", "S/D"), ss))
    story.append(_linea("2.7 Escala", form.get("escala", ""), ss))
    story.append(_linea("Tipo y ubicacion aproximada", form.get("escala_ubicacion", ""), ss))
    story.append(_linea("2.8 Mojon de referencia", form.get("mojon", ""), ss))
    story.append(_linea("Ubicacion aproximada", form.get("mojon_ubicacion", ""), ss))
    story.append(_linea("Cota", form.get("mojon_cota", ""), ss))
    story.append(_linea("2.9 Origen utilizado en la nivelacion", form.get("origen", ""), ss))

    story.append(PageBreak())

    # -------------------------
    # Pagina 2
    # -------------------------
    story.append(_pb("2.10 OBRAS DE TOMA", ss["Seccion"]))
    story.append(_linea("Ubicacion 1", form.get("toma_ubicacion", ""), ss))
    story.append(_linea("Tipo y diametro 1", form.get("toma_tipo", ""), ss))
    story.append(_linea("Cota de zampeado aguas abajo 1", form.get("toma_zampeado", ""), ss))

    story.append(_pb("2.11 SEGURIDAD DE LA REPRESA", ss["Seccion"]))
    story.append(_linea("(Indicar si la situacion observada hace temer por la seguridad de la obra)", form.get("seguridad", ""), ss))
    story.append(_p(f": {form.get('seguridad_comentarios', '')}", ss["Texto"]))

    story.append(_pb("2.12 VERTEDERO", ss["Seccion"]))
    story.append(_linea("Franquia (cota minima de corona - cota minima del umbral del vertedero)", auto.get("franquia", "S/D"), ss))
    story.append(_linea("Cota minima de corona utilizada", auto.get("cota_min_corona_franquia", "S/D"), ss))
    story.append(_linea("Cota minima de vertedero utilizada", auto.get("cota_min_vertedero_franquia", "S/D"), ss))
    story.append(_linea("Ubicacion", form.get("vertedero_ubicacion", ""), ss))
    story.append(_linea("Ancho del vertedero", auto.get("ancho_vertedero", "S/D"), ss))
    story.append(_p(f": Cotas del vertedero: {auto.get('cotas_vertedero', 'S/D')}", ss["Texto"]))
    story.append(_p(f": {form.get('vertedero_comentarios', '')}", ss["Texto"]))

    story.append(Spacer(1, 0.2 * cm))
    story.append(_p("<b>Comentarios</b>", ss["Texto"]))
    story.append(_p(
        f": La longitud del dique es de {auto.get('longitud', 'S/D')}, "
        f"con anchos de coronamiento entre {auto.get('ancho_min', 'S/D')} y {auto.get('ancho_max', 'S/D')}.",
        ss["Texto"]
    ))

    story.append(_pb("2.13 OBSERVACIONES:", ss["Seccion"]))
    story.append(_p(f": {form.get('observaciones', '')}", ss["Texto"]))

    story.append(_linea("2.14 Forma de acceso a la represa", form.get("acceso", ""), ss))
    story.append(_linea("2.15 Inspeccionada por", form.get("inspectores", ""), ss))

    story.append(PageBreak())

    # -------------------------
    # Pagina 3 - perfil
    # -------------------------
    story.append(_pb("3. Perfiles Longitudinales de Vertedero y Dique", ss["Seccion"]))
    story.append(Image(profile_path, width=16.2 * cm, height=7.2 * cm))
    story.append(_p("Figura 1. Perfil longitudinal generado a partir de los puntos seleccionados.", ss["Caption"]))

    if perfil_table is not None and len(perfil_table):
        tabla_perfil = perfil_table.copy()
        tabla_perfil = tabla_perfil.replace([np.inf, -np.inf], np.nan).fillna("")
        tabla_perfil = tabla_perfil.head(18)
        story.append(_tabla_pdf(tabla_perfil, font_size=7))
        if len(perfil_table) > 18:
            story.append(_p(f"Nota: se muestran las primeras 18 filas de {len(perfil_table)} puntos del perfil.", ss["NotaTabla"]))

    story.append(PageBreak())

    # -------------------------
    # Pagina 4 - MDT
    # -------------------------
    story.append(_pb("4. Modelo Digital de Terreno", ss["Seccion"]))
    story.append(Image(mdt_path, width=16.2 * cm, height=13.2 * cm))
    story.append(_p("Figura 2. Modelo Digital del Terreno generado mediante triangulacion TIN.", ss["Caption"]))

    story.append(PageBreak())

    # -------------------------
    # Pagina 5 - anexos
    # -------------------------
    story.append(_pb("5. Anexos", ss["Seccion"]))
    story.append(_pb("5.1 Anchos de coronamiento", ss["Seccion"]))

    if anchos_df is not None and len(anchos_df):
        resumen = pd.DataFrame({
            "Indicador": ["Anchos calculados", "Ancho minimo", "Ancho maximo", "Ancho promedio"],
            "Valor": [
                len(anchos_df),
                f"{anchos_df['Ancho corona (m)'].min():.3f} m",
                f"{anchos_df['Ancho corona (m)'].max():.3f} m",
                f"{anchos_df['Ancho corona (m)'].mean():.3f} m",
            ]
        })
        story.append(_tabla_pdf(resumen, font_size=9))
        story.append(Spacer(1, 0.35 * cm))
        story.append(_tabla_pdf(anchos_df.round(3), font_size=7))
    else:
        story.append(_p("No se calcularon anchos de coronamiento.", ss["Texto"]))

    story.append(PageBreak())

    # -------------------------
    # Pagina 6 - taludes
    # -------------------------
    story.append(_pb("5.2 Taludes", ss["Seccion"]))

    if talud_arriba_info:
        story.append(_pb("5.2.1 Talud aguas arriba", ss["Seccion"]))
        talud_arriba_df = pd.DataFrame({
            "Parametro": ["Punto superior", "Punto inferior", "Distancia horizontal H", "Diferencia vertical V", "Relacion"],
            "Valor": [
                talud_arriba_info["p_sup"],
                talud_arriba_info["p_inf"],
                f"{talud_arriba_info['H']:.3f} m",
                f"{talud_arriba_info['V']:.3f} m",
                f"1V : {talud_arriba_info['relacion']:.2f}H"
            ]
        })
        story.append(_tabla_pdf(talud_arriba_df, font_size=9))
        if incluir_talud_arriba and talud_arriba_path:
            story.append(Spacer(1, 0.35 * cm))
            story.append(Image(talud_arriba_path, width=14.5 * cm, height=8.0 * cm))
            story.append(_p("Figura. Esquema de calculo del talud aguas arriba.", ss["Caption"]))
        story.append(Spacer(1, 0.35 * cm))

    if talud_info:
        story.append(_pb("5.2.2 Talud aguas abajo", ss["Seccion"]))
        talud_df = pd.DataFrame({
            "Parametro": ["Punto superior", "Punto inferior", "Distancia horizontal H", "Diferencia vertical V", "Relacion"],
            "Valor": [
                talud_info["p_sup"],
                talud_info["p_inf"],
                f"{talud_info['H']:.3f} m",
                f"{talud_info['V']:.3f} m",
                f"1V : {talud_info['relacion']:.2f}H"
            ]
        })
        story.append(_tabla_pdf(talud_df, font_size=9))
        if incluir_talud_abajo and talud_path:
            story.append(Spacer(1, 0.35 * cm))
            story.append(Image(talud_path, width=14.5 * cm, height=8.0 * cm))
            story.append(_p("Figura. Esquema de calculo del talud aguas abajo.", ss["Caption"]))
    else:
        story.append(_p("No se calculo talud aguas abajo.", ss["Texto"]))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _guardar_figura_word(fig, nombre):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f"_{nombre}.png")
    fig.savefig(tmp.name, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return tmp.name

def _docx_heading(doc, text, size=12, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = True
    r.font.name = "Verdana"
    r.font.size = Pt(size)

def _docx_line(doc, label, value=""):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Verdana"
    r1.font.size = Pt(11)
    r2 = p.add_run(str(value))
    r2.font.name = "Verdana"
    r2.font.size = Pt(11)

def _docx_table(doc, df, max_rows=None):
    if df is None or len(df) == 0:
        doc.add_paragraph("Sin datos.")
        return
    dfx = df.copy()
    if max_rows is not None:
        dfx = dfx.head(max_rows)
    table = doc.add_table(rows=1, cols=len(dfx.columns))
    table.style = "Table Grid"
    for i, col in enumerate(dfx.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in dfx.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)

def build_word(form, auto, data, fig_mdt, fig_profile, anchos_df, talud_info, fig_talud, perfil_table,
               incluir_talud_abajo=True, talud_arriba_info=None, fig_talud_arriba=None, incluir_talud_arriba=False):
    buffer = io.BytesIO()
    doc = Document()
    doc.styles["Normal"].font.name = "Verdana"
    doc.styles["Normal"].font.size = Pt(11)

    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.0)

    mdt_path = _guardar_figura_word(fig_mdt, "mdt_word")
    perfil_path = _guardar_figura_word(fig_profile, "perfil_word")
    talud_path = _guardar_figura_word(fig_talud, "talud_word") if fig_talud is not None else None
    talud_arriba_path = _guardar_figura_word(fig_talud_arriba, "talud_arriba_word") if fig_talud_arriba is not None else None

    _docx_heading(doc, "PLANILLA INSPECCION", size=13, center=True)

    _docx_heading(doc, "1. DATOS GENERALES")
    _docx_line(doc, "1.1 N° de represa", form.get("n_represa", ""))
    _docx_line(doc, "1.2 Fecha inspeccion", form.get("fecha", ""))
    _docx_line(doc, "1.3 Curso embalsado", form.get("curso", ""))
    _docx_line(doc, "1.4 Nombre del propietario", form.get("propietario", ""))
    _docx_line(doc, "1.5 Padrones represa", form.get("padrones", ""))
    _docx_line(doc, "Seccion Judicial", form.get("seccion", ""))
    _docx_line(doc, "Departamento", form.get("departamento", ""))

    _docx_heading(doc, "2. DATOS REPRESA")
    _docx_line(doc, "2.1 Ancho minimo de coronamiento", auto.get("ancho_min", "S/D"))
    _docx_line(doc, "Altura del terraplen", auto.get("altura_max", "S/D"))
    _docx_line(doc, "Ancho de coronamiento", auto.get("ancho_prom", "S/D"))

    _docx_heading(doc, "2.3 TALUD AGUAS ARRIBA")
    doc.add_paragraph(str(form.get("talud_arriba", "")) + " Pendiente calculada = " + str(auto.get("talud_arriba_calc", "S/D")))
    _docx_heading(doc, "2.4 TALUD AGUAS ABAJO")
    doc.add_paragraph(str(form.get("talud_abajo", "")) + " Pendiente calculada = " + str(auto.get("talud", "S/D")))

    _docx_line(doc, "2.5 Cota del Pelo de agua", form.get("pelo_agua", ""))
    _docx_line(doc, "2.6 Cota del Coronamiento", auto.get("cota_corona", "S/D"))
    _docx_line(doc, "2.7 Escala", form.get("escala", ""))
    _docx_line(doc, "Tipo y ubicacion aproximada", form.get("escala_ubicacion", ""))
    _docx_line(doc, "2.8 Mojon de referencia", form.get("mojon", ""))
    _docx_line(doc, "Ubicacion aproximada", form.get("mojon_ubicacion", ""))
    _docx_line(doc, "Cota", form.get("mojon_cota", ""))
    _docx_line(doc, "2.9 Origen utilizado en la nivelacion", form.get("origen", ""))

    doc.add_page_break()
    _docx_heading(doc, "2.10 OBRAS DE TOMA")
    _docx_line(doc, "Ubicacion", form.get("toma_ubicacion", ""))
    _docx_line(doc, "Tipo y diametro", form.get("toma_tipo", ""))
    _docx_line(doc, "Cota de zampeado aguas abajo", form.get("toma_zampeado", ""))

    _docx_heading(doc, "2.11 SEGURIDAD DE LA REPRESA")
    _docx_line(doc, "Situacion observada hace temer por la seguridad de la obra", form.get("seguridad", ""))
    doc.add_paragraph(str(form.get("seguridad_comentarios", "")))

    _docx_heading(doc, "2.12 VERTEDERO")
    _docx_line(doc, "Franquia", auto.get("franquia", "S/D"))
    _docx_line(doc, "Cota minima de corona utilizada", auto.get("cota_min_corona_franquia", "S/D"))
    _docx_line(doc, "Cota minima de vertedero utilizada", auto.get("cota_min_vertedero_franquia", "S/D"))
    _docx_line(doc, "Ubicacion", form.get("vertedero_ubicacion", ""))
    _docx_line(doc, "Ancho del vertedero", auto.get("ancho_vertedero", "S/D"))
    _docx_line(doc, "Cotas del vertedero", auto.get("cotas_vertedero", "S/D"))
    doc.add_paragraph(str(form.get("vertedero_comentarios", "")))

    _docx_heading(doc, "Comentarios")
    doc.add_paragraph(
        f"La longitud del dique es de {auto.get('longitud', 'S/D')}, "
        f"con anchos de coronamiento entre {auto.get('ancho_min', 'S/D')} y {auto.get('ancho_max', 'S/D')}."
    )

    _docx_heading(doc, "2.13 OBSERVACIONES")
    doc.add_paragraph(str(form.get("observaciones", "")))
    _docx_line(doc, "2.14 Forma de acceso a la represa", form.get("acceso", ""))
    _docx_line(doc, "2.15 Inspeccionada por", form.get("inspectores", ""))

    doc.add_page_break()
    _docx_heading(doc, "3. Perfiles Longitudinales de Vertedero y Dique")
    doc.add_picture(perfil_path, width=Cm(16.2))
    doc.add_paragraph("Figura 1. Perfil longitudinal generado a partir de los puntos seleccionados.")
    if perfil_table is not None and len(perfil_table):
        tabla_perfil = perfil_table.copy().replace([np.inf, -np.inf], np.nan).fillna("")
        _docx_table(doc, tabla_perfil, max_rows=18)
        if len(perfil_table) > 18:
            doc.add_paragraph(f"Nota: se muestran las primeras 18 filas de {len(perfil_table)} puntos del perfil.")

    doc.add_page_break()
    _docx_heading(doc, "4. Modelo Digital de Terreno")
    doc.add_picture(mdt_path, width=Cm(16.2))
    doc.add_paragraph("Figura 2. Modelo Digital del Terreno generado mediante triangulacion TIN.")

    doc.add_page_break()
    _docx_heading(doc, "5. Anexos")
    _docx_heading(doc, "5.1 Anchos de coronamiento")
    if anchos_df is not None and len(anchos_df):
        resumen = pd.DataFrame({
            "Indicador": ["Anchos calculados", "Ancho minimo", "Ancho maximo", "Ancho promedio"],
            "Valor": [
                len(anchos_df),
                f"{anchos_df['Ancho corona (m)'].min():.3f} m",
                f"{anchos_df['Ancho corona (m)'].max():.3f} m",
                f"{anchos_df['Ancho corona (m)'].mean():.3f} m",
            ]
        })
        _docx_table(doc, resumen)
        doc.add_paragraph("")
        _docx_table(doc, anchos_df.round(3))
    else:
        doc.add_paragraph("No se calcularon anchos de coronamiento.")

    doc.add_page_break()
    _docx_heading(doc, "5.2 Taludes")
    if talud_arriba_info:
        _docx_heading(doc, "5.2.1 Talud aguas arriba")
        _docx_table(doc, pd.DataFrame({
            "Parametro": ["Punto superior", "Punto inferior", "Distancia horizontal H", "Diferencia vertical V", "Relacion"],
            "Valor": [
                talud_arriba_info["p_sup"],
                talud_arriba_info["p_inf"],
                f"{talud_arriba_info['H']:.3f} m",
                f"{talud_arriba_info['V']:.3f} m",
                f"1V : {talud_arriba_info['relacion']:.2f}H"
            ]
        }))
        if incluir_talud_arriba and talud_arriba_path:
            doc.add_picture(talud_arriba_path, width=Cm(14.5))

    if talud_info:
        _docx_heading(doc, "5.2.2 Talud aguas abajo")
        _docx_table(doc, pd.DataFrame({
            "Parametro": ["Punto superior", "Punto inferior", "Distancia horizontal H", "Diferencia vertical V", "Relacion"],
            "Valor": [
                talud_info["p_sup"],
                talud_info["p_inf"],
                f"{talud_info['H']:.3f} m",
                f"{talud_info['V']:.3f} m",
                f"1V : {talud_info['relacion']:.2f}H"
            ]
        }))
        if incluir_talud_abajo and talud_path:
            doc.add_picture(talud_path, width=Cm(14.5))

    doc.save(buffer)
    buffer.seek(0)
    return buffer


if uploaded:
    df = read_file(uploaded)
    st.subheader("Vista previa")
    st.dataframe(df.head(30), use_container_width=True)

    cols = list(df.columns)
    default_num = cols.index("Numero") if "Numero" in cols else 0
    default_x = cols.index("X") if "X" in cols else min(1, len(cols)-1)
    default_y = cols.index("Y") if "Y" in cols else min(2, len(cols)-1)
    default_z = cols.index("Z") if "Z" in cols else min(3, len(cols)-1)
    if "Descripción" in cols:
        default_desc = cols.index("Descripción")
    elif "Descripcion" in cols:
        default_desc = cols.index("Descripcion")
    else:
        default_desc = min(4, len(cols)-1)

    st.subheader("Seleccionar columnas")
    c1, c2, c3, c4, c5 = st.columns(5)
    num_col = c1.selectbox("Numero / ID", cols, index=default_num)
    x_col = c2.selectbox("X / Este", cols, index=default_x)
    y_col = c3.selectbox("Y / Norte", cols, index=default_y)
    z_col = c4.selectbox("Z / Cota", cols, index=default_z)
    desc_col = c5.selectbox("Descripción", cols, index=default_desc)

    data = df[[num_col, x_col, y_col, z_col, desc_col]].copy()
    data.columns = ["Numero", "X", "Y", "Z", "Descripcion"]
    data["X"] = pd.to_numeric(data["X"], errors="coerce")
    data["Y"] = pd.to_numeric(data["Y"], errors="coerce")
    data["Z"] = pd.to_numeric(data["Z"], errors="coerce")
    data = data.dropna(subset=["X", "Y", "Z"]).reset_index(drop=True)
    data["Descripcion_norm"] = data["Descripcion"].apply(norm_text)

    if len(data) < 3:
        st.error("Para generar MDT necesitás al menos 3 puntos.")
        st.stop()

    descripciones = sorted(data["Descripcion_norm"].dropna().unique().tolist())

    st.subheader("Resumen general")
    a, b, c, d = st.columns(4)
    a.metric("Puntos", len(data))
    b.metric("Cota mínima", f"{data['Z'].min():.3f} m")
    c.metric("Cota máxima", f"{data['Z'].max():.3f} m")
    d.metric("Desnivel", f"{data['Z'].max() - data['Z'].min():.3f} m")

    st.divider()
    st.header("Formulario de inspección")

    with st.expander("1. Datos generales", expanded=True):
        g1, g2 = st.columns(2)
        n_represa = g1.text_input("1.1 N° de represa", "")
        fecha = g2.date_input("1.2 Fecha inspección")
        curso = st.text_input("1.3 Curso embalsado", "")
        propietario = st.text_input("1.4 Nombre del propietario", "")
        p1, p2, p3 = st.columns(3)
        padrones = p1.text_input("1.5 Padrones represa", "")
        seccion = p2.text_input("Sección Judicial", "")
        departamento = p3.text_input("Departamento", "")

    with st.expander("2. Datos descriptivos de represa", expanded=True):
        talud_arriba = st.text_area("2.3 Talud aguas arriba", "Describir estado actual: pendiente, enrocado, erosiones.")
        talud_abajo_txt = st.text_area("2.4 Talud aguas abajo", "Describir estado actual: pendiente, filtraciones, empastado, erosiones.")
        pelo_agua = st.text_input("2.5 Cota del pelo de agua", "")
        escala = st.selectbox("2.7 Escala", ["NO", "SI"])
        escala_ubicacion = st.text_input("Tipo y ubicación aproximada de escala", "")
        mojon = st.selectbox("2.8 Mojón de referencia", ["SI", "NO"])
        mojon_ubicacion = st.text_input("Ubicación aproximada del mojón", "")
        mojon_cota = st.text_input("Cota del mojón", "")
        origen = st.text_input("2.9 Origen utilizado en la nivelación", "Mojón")

    with st.expander("2.10 a 2.15 Obras, seguridad y observaciones", expanded=True):
        toma_ubicacion = st.text_input("2.10 Ubicación obra de toma", "")
        toma_tipo = st.text_input("Tipo y diámetro", "")
        toma_zampeado = st.text_input("Cota de zampeado aguas abajo", "")
        seguridad = st.selectbox("2.11 Seguridad de la represa", ["NO", "SI", "ATENCIÓN"])
        seguridad_comentarios = st.text_area("Comentarios seguridad", "")
        vertedero_ubicacion = st.text_input("2.12 Ubicación vertedero", "")
        vertedero_ancho_manual = st.text_input("Ancho del vertedero (manual)", "")
        vertedero_cota_min_manual = st.text_input("Cota mínima real del umbral del vertedero (manual)", "")
        vertedero_cota_max_manual = st.text_input("Cota máxima real del umbral del vertedero (manual)", "")
        vertedero_comentarios = st.text_area("Comentarios vertedero", "")
        observaciones = st.text_area("2.13 Observaciones", "")
        acceso = st.text_area("2.14 Forma de acceso a la represa", "")
        inspectores = st.text_input("2.15 Inspeccionada por", "")

    form = {
        "n_represa": n_represa, "fecha": fecha.strftime("%d/%m/%Y"), "curso": curso, "propietario": propietario,
        "padrones": padrones, "seccion": seccion, "departamento": departamento,
        "talud_arriba": talud_arriba, "talud_abajo": talud_abajo_txt, "pelo_agua": pelo_agua,
        "escala": escala, "escala_ubicacion": escala_ubicacion, "mojon": mojon, "mojon_ubicacion": mojon_ubicacion,
        "mojon_cota": mojon_cota, "origen": origen, "toma_ubicacion": toma_ubicacion, "toma_tipo": toma_tipo,
        "toma_zampeado": toma_zampeado, "seguridad": seguridad, "seguridad_comentarios": seguridad_comentarios,
        "vertedero_ubicacion": vertedero_ubicacion,
        "vertedero_ancho_manual": vertedero_ancho_manual,
        "vertedero_cota_min_manual": vertedero_cota_min_manual,
        "vertedero_cota_max_manual": vertedero_cota_max_manual,
        "vertedero_comentarios": vertedero_comentarios,
        "observaciones": observaciones, "acceso": acceso, "inspectores": inspectores
    }

    st.divider()
    st.header("Cálculos automáticos")

    st.info("Nota: las cotas y el ancho del vertedero para el informe se cargan manualmente. No se toman de los puntos inicio/final del vertedero.")

    fc1, fc2, fc3 = st.columns(3)
    corona_sel = fc1.selectbox("Descripción de corona", descripciones, index=descripciones.index("corona") if "corona" in descripciones else 0)
    vertedero_sel = fc2.selectbox("Descripción de vertedero", descripciones, index=descripciones.index("vertedero") if "vertedero" in descripciones else (1 if len(descripciones)>1 else 0))
    base_sel = fc3.selectbox("Descripción de base", descripciones, index=descripciones.index("base") if "base" in descripciones else 0)

    corona = data[data["Descripcion_norm"] == corona_sel].copy()
    vertedero = data[data["Descripcion_norm"] == vertedero_sel].copy()
    base_points = data[data["Descripcion_norm"] == base_sel].copy()
    anchos_df = calcular_anchos_pares(base_points) if len(base_points) >= 2 else pd.DataFrame()

    auto = {}
    if len(corona):
        auto["cota_corona"] = f"Entre {corona['Z'].min():.3f} m y {corona['Z'].max():.3f} m"
        corona_prof = calcular_distancia_acumulada(corona.sort_index())
        auto["longitud"] = f"{corona_prof['Distancia acumulada'].iloc[-1]:.2f} m"
    else:
        auto["cota_corona"] = "S/D"
        auto["longitud"] = "S/D"

    # Datos manuales del vertedero:
    # No se usan automáticamente las cotas de los puntos "vertedero",
    # porque el primer y último punto pueden ser inicio/final y no el umbral real.
    cota_min_vertedero_manual = parse_float_manual(form.get("vertedero_cota_min_manual", ""))
    cota_max_vertedero_manual = parse_float_manual(form.get("vertedero_cota_max_manual", ""))

    auto["ancho_vertedero"] = form.get("vertedero_ancho_manual", "").strip() or "S/D"

    if cota_min_vertedero_manual is not None and cota_max_vertedero_manual is not None:
        auto["cotas_vertedero"] = f"Entre {cota_min_vertedero_manual:.3f} m y {cota_max_vertedero_manual:.3f} m"
    elif cota_min_vertedero_manual is not None:
        auto["cotas_vertedero"] = f"Cota mínima del umbral: {cota_min_vertedero_manual:.3f} m"
    elif cota_max_vertedero_manual is not None:
        auto["cotas_vertedero"] = f"Cota máxima del umbral: {cota_max_vertedero_manual:.3f} m"
    else:
        auto["cotas_vertedero"] = "S/D"

    # Franquía = cota mínima de corona - cota mínima real del umbral del vertedero.
    # La cota de corona sale de los puntos seleccionados como "corona".
    # La cota del vertedero se ingresa manualmente porque los puntos inicio/final
    # pueden no representar el umbral real.
    if len(corona) and cota_min_vertedero_manual is not None:
        cota_min_corona = float(corona["Z"].min())
        franquia = cota_min_corona - cota_min_vertedero_manual
        auto["franquia"] = f"{franquia:.3f} m"
        auto["cota_min_corona_franquia"] = f"{cota_min_corona:.3f} m"
        auto["cota_min_vertedero_franquia"] = f"{cota_min_vertedero_manual:.3f} m"
    else:
        auto["franquia"] = "S/D"
        auto["cota_min_corona_franquia"] = "S/D"
        auto["cota_min_vertedero_franquia"] = "S/D"

    if len(anchos_df):
        auto["ancho_min"] = f"{anchos_df['Ancho corona (m)'].min():.3f} m"
        auto["ancho_prom"] = f"{anchos_df['Ancho corona (m)'].mean():.3f} m"
        auto["ancho_max"] = f"{anchos_df['Ancho corona (m)'].max():.3f} m"
    else:
        auto["ancho_min"] = auto["ancho_prom"] = auto["ancho_max"] = "S/D"

    auto["altura_max"] = f"{data['Z'].max() - data['Z'].min():.3f} m"

    r1, r2, r3, r4 = st.columns(4)
    r1.metric("Longitud corona", auto["longitud"])
    r2.metric("Franquía", auto["franquia"])
    r3.metric("Ancho mínimo", auto["ancho_min"])
    r4.metric("Altura / desnivel máx.", auto["altura_max"])

    if len(anchos_df):
        st.subheader("Tabla de anchos")
        st.dataframe(anchos_df.round(3), use_container_width=True)

    st.divider()
    st.header("MDT")
    m1, m2, m3, m4 = st.columns(4)
    cmap = m1.selectbox("Rampa de colores", ["terrain", "viridis", "plasma", "turbo", "jet", "gist_earth"])
    mostrar_curvas = m2.checkbox("Mostrar curvas", True)
    intervalo = m3.number_input("Intervalo curvas (m)", value=0.20, step=0.05, min_value=0.01)
    curvas_etiquetas = m4.checkbox("Etiquetar curvas", True)
    fig_mdt = create_mdt_figure(data, cmap, mostrar_curvas, intervalo, curvas_etiquetas)
    st.pyplot(fig_mdt)


    st.divider()
    st.header("Talud aguas arriba manual")
    st.write("Seleccioná el punto superior y el punto inferior del talud aguas arriba. La app calcula H, V y la relación 1V : nH.")

    opciones_puntos = [f"{row['Numero']} | {row['Descripcion']} | Z={row['Z']:.3f}" for _, row in data.iterrows()]

    ta1, ta2 = st.columns(2)
    punto_arr_sup_txt = ta1.selectbox(
        "Punto superior - talud aguas arriba",
        opciones_puntos,
        index=0,
        key="talud_arriba_superior"
    )
    punto_arr_inf_txt = ta2.selectbox(
        "Punto inferior - talud aguas arriba",
        opciones_puntos,
        index=min(1, len(opciones_puntos)-1),
        key="talud_arriba_inferior"
    )

    p_arr_sup = data.iloc[opciones_puntos.index(punto_arr_sup_txt)]
    p_arr_inf = data.iloc[opciones_puntos.index(punto_arr_inf_txt)]

    H_arriba = math.sqrt((p_arr_inf["X"] - p_arr_sup["X"])**2 + (p_arr_inf["Y"] - p_arr_sup["Y"])**2)
    V_arriba = abs(p_arr_inf["Z"] - p_arr_sup["Z"])

    talud_arriba_info = None
    fig_talud_arriba = None

    if V_arriba == 0:
        st.error("La diferencia vertical del talud aguas arriba es 0. No se puede calcular.")
        auto["talud_arriba_calc"] = "S/D"
    else:
        relacion_arriba = H_arriba / V_arriba
        auto["talud_arriba_calc"] = f"1V : {relacion_arriba:.2f}H"

        talud_arriba_info = {
            "p_sup": p_arr_sup["Numero"],
            "p_inf": p_arr_inf["Numero"],
            "H": H_arriba,
            "V": V_arriba,
            "relacion": relacion_arriba
        }

        taq1, taq2, taq3 = st.columns(3)
        taq1.metric("H aguas arriba", f"{H_arriba:.3f} m")
        taq2.metric("V aguas arriba", f"{V_arriba:.3f} m")
        taq3.metric("Talud aguas arriba", auto["talud_arriba_calc"])

        fig_talud_arriba = plot_talud(H_arriba, V_arriba, p_arr_sup["Numero"], p_arr_inf["Numero"], relacion_arriba)
        st.pyplot(fig_talud_arriba)

    st.divider()
    st.header("Talud aguas abajo manual")
    opciones_puntos = [f"{row['Numero']} | {row['Descripcion']} | Z={row['Z']:.3f}" for _, row in data.iterrows()]
    t1, t2 = st.columns(2)
    punto_sup_txt = t1.selectbox("Punto superior", opciones_puntos, index=0)
    punto_inf_txt = t2.selectbox("Punto inferior", opciones_puntos, index=min(1, len(opciones_puntos)-1))
    ps = data.iloc[opciones_puntos.index(punto_sup_txt)]
    pi = data.iloc[opciones_puntos.index(punto_inf_txt)]
    H = math.sqrt((pi["X"] - ps["X"])**2 + (pi["Y"] - ps["Y"])**2)
    V = abs(pi["Z"] - ps["Z"])
    talud_info = None
    fig_talud = None
    if V == 0:
        st.error("La diferencia vertical es 0. No se puede calcular el talud.")
        auto["talud"] = "S/D"
    else:
        relacion = H / V
        auto["talud"] = f"1V : {relacion:.2f}H"
        talud_info = {"p_sup": ps["Numero"], "p_inf": pi["Numero"], "H": H, "V": V, "relacion": relacion}
        q1, q2, q3 = st.columns(3)
        q1.metric("H", f"{H:.3f} m")
        q2.metric("V", f"{V:.3f} m")
        q3.metric("Talud", auto["talud"])
        fig_talud = plot_talud(H, V, ps["Numero"], pi["Numero"], relacion)
        st.pyplot(fig_talud)


    st.divider()
    st.header("Altura máxima en sección de máxima altura")
    st.write("Seleccioná manualmente el punto superior de coronamiento y el punto inferior de la sección de máxima altura.")

    h1, h2 = st.columns(2)
    punto_alt_sup_txt = h1.selectbox(
        "Punto superior para altura máxima",
        opciones_puntos,
        index=0,
        key="altura_punto_superior"
    )
    punto_alt_inf_txt = h2.selectbox(
        "Punto inferior para altura máxima",
        opciones_puntos,
        index=min(1, len(opciones_puntos)-1),
        key="altura_punto_inferior"
    )

    p_alt_sup = data.iloc[opciones_puntos.index(punto_alt_sup_txt)]
    p_alt_inf = data.iloc[opciones_puntos.index(punto_alt_inf_txt)]

    altura_max_manual = abs(p_alt_sup["Z"] - p_alt_inf["Z"])
    distancia_horizontal_altura = math.sqrt((p_alt_inf["X"] - p_alt_sup["X"])**2 + (p_alt_inf["Y"] - p_alt_sup["Y"])**2)

    auto["altura_max"] = f"{altura_max_manual:.3f} m"

    ah1, ah2, ah3 = st.columns(3)
    ah1.metric("Altura máxima manual", f"{altura_max_manual:.3f} m")
    ah2.metric("Distancia horizontal entre puntos", f"{distancia_horizontal_altura:.3f} m")
    ah3.metric("Puntos usados", f"{p_alt_sup['Numero']} - {p_alt_inf['Numero']}")

    altura_max_df = pd.DataFrame({
        "Parámetro": [
            "Punto superior",
            "Cota superior",
            "Punto inferior",
            "Cota inferior",
            "Altura máxima",
            "Distancia horizontal"
        ],
        "Valor": [
            str(p_alt_sup["Numero"]),
            f"{p_alt_sup['Z']:.3f} m",
            str(p_alt_inf["Numero"]),
            f"{p_alt_inf['Z']:.3f} m",
            f"{altura_max_manual:.3f} m",
            f"{distancia_horizontal_altura:.3f} m"
        ]
    })

    st.subheader("Resumen sección de máxima altura")
    st.dataframe(altura_max_df, use_container_width=True)

    st.divider()
    st.header("Perfil longitudinal")
    perfil_sel = st.multiselect("Descripción/es para perfil", descripciones, default=[corona_sel])
    perfil = data[data["Descripcion_norm"].isin(perfil_sel)].copy().sort_index().reset_index(drop=True)
    fig_profile = None
    perfil_table = None

    if len(perfil) >= 2:
        perfil_temp = calcular_distancia_acumulada(perfil)

        mp1, mp2, mp3, mp4 = st.columns(4)
        mp1.metric("Puntos perfil", len(perfil))
        mp2.metric("Longitud perfil", f"{perfil_temp['Distancia acumulada'].iloc[-1]:.2f} m")
        mp3.metric("Cota mínima", f"{perfil['Z'].min():.3f} m")
        mp4.metric("Cota máxima", f"{perfil['Z'].max():.3f} m")

        st.subheader("Opciones de escala y visualización del perfil")

        y_min_default = math.floor(float(perfil["Z"].min()) * 2) / 2
        y_max_default = math.ceil(float(perfil["Z"].max()) * 2) / 2

        ep1, ep2, ep3, ep4 = st.columns(4)
        y_min_perfil = ep1.number_input("Cota mínima del gráfico", value=float(y_min_default), step=0.1)
        y_max_perfil = ep2.number_input("Cota máxima del gráfico", value=float(y_max_default), step=0.1)
        mostrar_marcadores = ep3.checkbox("Mostrar marcadores", value=True)
        sombreado_perfil = ep4.checkbox("Sombreado bajo perfil", value=True)

        st.subheader("Carteles de mínima y máxima")
        cp1, cp2 = st.columns(2)
        mostrar_min = cp1.checkbox("Mostrar cartel mínimo", value=True)
        mostrar_max = cp2.checkbox("Mostrar cartel máximo", value=True)
        st.caption("X positivo mueve el cartel a la derecha. Y positivo lo mueve hacia arriba.")

        cmin1, cmin2, cmax1, cmax2 = st.columns(4)
        min_dx = cmin1.slider("Mínimo mover X", -180, 180, 15, 5)
        min_dy = cmin2.slider("Mínimo mover Y", -120, 120, -35, 5)
        max_dx = cmax1.slider("Máximo mover X", -180, 180, -80, 5)
        max_dy = cmax2.slider("Máximo mover Y", -120, 120, 25, 5)

        fig_profile, perfil_proc = create_profile_figure(
            perfil,
            ", ".join(perfil_sel),
            y_min_manual=y_min_perfil,
            y_max_manual=y_max_perfil,
            mostrar_marcadores=mostrar_marcadores,
            sombreado=sombreado_perfil,
            mostrar_min=mostrar_min,
            mostrar_max=mostrar_max,
            min_dx=min_dx,
            min_dy=min_dy,
            max_dx=max_dx,
            max_dy=max_dy
        )

        perfil_table = perfil_proc[["Numero", "Descripcion", "Progresiva", "Distancia acumulada", "Z"]].rename(columns={"Z": "Cota"})
        st.pyplot(fig_profile)
        st.dataframe(perfil_table.round(3), use_container_width=True)
    else:
        st.info("Seleccioná al menos 2 puntos para el perfil.")

    st.divider()
    st.header("Opciones del informe PDF")
    op_pdf1, op_pdf2 = st.columns(2)
    incluir_grafico_talud_arriba = op_pdf1.checkbox("Incluir gráfico de Talud Aguas Arriba en el PDF", value=True)
    incluir_grafico_talud_abajo = op_pdf2.checkbox("Incluir gráfico de Talud Aguas Abajo en el PDF", value=True)

    st.divider()
    st.header("Generar PDF estilo planilla de inspección")
    if st.button("📄 Generar informe PDF"):
        if fig_profile is None:
            st.error("Primero generá un perfil válido.")
        elif fig_talud is None:
            st.error("Primero calculá un talud válido.")
        else:
            pdf_buffer = build_pdf(
                form, auto, data,
                create_mdt_figure(data, cmap, mostrar_curvas, intervalo, curvas_etiquetas, for_pdf=True),
                create_profile_figure(
                    perfil,
                    ", ".join(perfil_sel),
                    for_pdf=True,
                    y_min_manual=y_min_perfil,
                    y_max_manual=y_max_perfil,
                    mostrar_marcadores=mostrar_marcadores,
                    sombreado=sombreado_perfil,
                    mostrar_min=mostrar_min,
                    mostrar_max=mostrar_max,
                    min_dx=min_dx,
                    min_dy=min_dy,
                    max_dx=max_dx,
                    max_dy=max_dy
                )[0],
                anchos_df,
                talud_info,
                plot_talud(talud_info["H"], talud_info["V"], talud_info["p_sup"], talud_info["p_inf"], talud_info["relacion"], for_pdf=True),
                perfil_table,
                incluir_talud_abajo=incluir_grafico_talud_abajo,
                talud_arriba_info=talud_arriba_info if "talud_arriba_info" in locals() else None,
                fig_talud_arriba=plot_talud(talud_arriba_info["H"], talud_arriba_info["V"], talud_arriba_info["p_sup"], talud_arriba_info["p_inf"], talud_arriba_info["relacion"], for_pdf=True) if "talud_arriba_info" in locals() and talud_arriba_info else None,
                incluir_talud_arriba=incluir_grafico_talud_arriba
            )
            st.download_button("Descargar PDF", pdf_buffer, "Planilla_Inspeccion_Represa.pdf", "application/pdf")

            word_buffer = build_word(
                form, auto, data,
                create_mdt_figure(data, cmap, mostrar_curvas, intervalo, curvas_etiquetas, for_pdf=True),
                create_profile_figure(
                    perfil,
                    ", ".join(perfil_sel),
                    for_pdf=True,
                    y_min_manual=y_min_perfil,
                    y_max_manual=y_max_perfil,
                    mostrar_marcadores=mostrar_marcadores,
                    sombreado=sombreado_perfil,
                    mostrar_min=mostrar_min,
                    mostrar_max=mostrar_max,
                    min_dx=min_dx,
                    min_dy=min_dy,
                    max_dx=max_dx,
                    max_dy=max_dy
                )[0],
                anchos_df,
                talud_info,
                plot_talud(talud_info["H"], talud_info["V"], talud_info["p_sup"], talud_info["p_inf"], talud_info["relacion"], for_pdf=True),
                perfil_table,
                incluir_talud_abajo=incluir_grafico_talud_abajo,
                talud_arriba_info=talud_arriba_info if "talud_arriba_info" in locals() else None,
                fig_talud_arriba=plot_talud(talud_arriba_info["H"], talud_arriba_info["V"], talud_arriba_info["p_sup"], talud_arriba_info["p_inf"], talud_arriba_info["relacion"], for_pdf=True) if "talud_arriba_info" in locals() and talud_arriba_info else None,
                incluir_talud_arriba=incluir_grafico_talud_arriba
            )
            st.download_button("Descargar Word", word_buffer, "Planilla_Inspeccion_Represa.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Subí tu archivo para empezar.")